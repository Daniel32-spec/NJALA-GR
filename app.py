import streamlit as st
import pandas as pd
import joblib
from sklearn.ensemble import RandomForestClassifier
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize session state for results history
if 'results_history' not in st.session_state:
    st.session_state.results_history = []

# Load and train model
@st.cache_resource
def train_model():
    try:
        df = pd.read_csv('njala_student_data.csv')
        df['Result'] = df['Result'].map({'Pass': 1, 'Fail': 0})
        X = df[['GPA', 'Credit_Hours', 'Year_Average']]
        y = df['Result']
        model = RandomForestClassifier(random_state=42)
        model.fit(X, y)
        return model
    except Exception as e:
        st.error(f"Error loading model: {str(e)}")
        return None

model = train_model()

# Function to convert score to letter grade
def score_to_letter_grade(score):
    try:
        score = float(score)
        if score >= 75:
            return "A"
        elif score >= 65:
            return "B"
        elif score >= 50:
            return "C"
        elif score >= 40:
            return "D"
        elif score >= 30:
            return "E"
        else:
            return "F"
    except (ValueError, TypeError):
        return None

# Function to determine pass/fail
def score_to_pass_fail(score):
    try:
        score = float(score)
        return "Pass" if score >= 50 else "Fail"
    except (ValueError, TypeError):
        return None

# Function to convert grades to GPA
def grade_to_gpa(grade):
    grade_map = {'A': 4.0, 'B': 3.0, 'C': 2.0, 'D': 1.0, 'E': 0.7, 'F': 0.0}
    try:
        grade = float(grade)
        if 0 <= grade <= 100:
            letter_grade = score_to_letter_grade(grade)
            return grade_map.get(letter_grade, None)
        return None
    except (ValueError, TypeError):
        grade = str(grade).strip().upper()
        return grade_map.get(grade, None)

# Function to create Word document with formatted table
def create_word_doc(title, data):
    doc = Document()
    doc.add_heading(title, 0)
    
    if isinstance(data, pd.DataFrame):
        table = doc.add_table(rows=data.shape[0] + 1, cols=len(data.columns))
        table.style = 'Table Grid'
        
        for j, col in enumerate(data.columns):
            cell = table.cell(0, j)
            cell.text = col
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row in data.iterrows():
            for j, val in enumerate(row):
                table.cell(i + 1, j).text = str(val)
        
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.0)
    
    else:
        doc.add_paragraph(str(data))
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# Function to create Excel file
def create_excel_file(data):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        data.to_excel(writer, index=False, sheet_name='Results')
    return output.getvalue()

# Sidebar navigation
st.sidebar.title("Njala University Tools")
page = st.sidebar.radio("Select Tool", ["Predictor", "Module Grades", "Semester Grades", "Yearly Grades"])

# Main app title
st.title("🎓 Njala University Academic Tools")
st.markdown("Select a tool to predict pass/fail or calculate grades. All sections support Excel input and Excel/Word output with tables for all variables.")

# Predictor Page
if page == "Predictor":
    st.subheader("Pass/Fail Predictor")
    st.markdown("""
        Predict pass/fail using manual input or an Excel file (columns: SN, NAME, ID, GRADE).
        Download results as Excel or Word with tables for all variables.
    """)

    tab1, tab2 = st.tabs(["Manual Input", "Upload Excel File"])

    with tab1:
        with st.form("prediction_form"):
            col1, col2, col3 = st.columns(3)
            with col1:
                module_name = st.text_input("📘 Module Name", placeholder="e.g., MATH202")
            with col2:
                module_credit = st.number_input("📐 Credit Hours", min_value=1, max_value=6, step=1, value=3)
            with col3:
                gpa = st.slider("GPA", min_value=1.0, max_value=4.0, step=0.01, value=2.5)
                year_avg = st.slider("Year Average", min_value=1.0, max_value=4.0, step=0.01, value=2.5)
            submitted = st.form_submit_button("Predict")
            clear = st.form_submit_button("Clear")

        def validate_inputs(module_name, gpa, module_credit, year_avg):
            if not module_name.strip():
                return False, "Module name cannot be empty"
            if not re.match(r'^[A-Z0-9]+$', module_name.strip()):
                return False, "Module name should contain only uppercase letters and numbers"
            if gpa < 1.0 or year_avg < 1.0:
                return False, "GPA and Year Average must be at least 1.0"
            return True, ""

        if submitted and model is not None:
            is_valid, error_message = validate_inputs(module_name, gpa, module_credit, year_avg)
            if is_valid:
                prediction = model.predict([[gpa, module_credit, year_avg]])[0]
                result = "Pass" if prediction == 1 else "Fail"
                st.session_state.results_history.append({
                    "Module": module_name,
                    "GPA": gpa,
                    "Credit Hours": module_credit,
                    "Year Average": year_avg,
                    "Result": result
                })
                st.success(f"Module: {module_name} | Predicted Result: {result}")
                output_data = pd.DataFrame([{
                    "Module": module_name,
                    "GPA": gpa,
                    "Credit Hours": module_credit,
                    "Year Average": year_avg,
                    "Result": result
                }])
                output_format = st.selectbox("Output format", ["Excel", "Word"], key="predictor_manual_out")
                if output_format == "Excel":
                    st.download_button("Download Excel", create_excel_file(output_data), "predictor.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                else:
                    st.download_button("Download Word", create_word_doc("Predictor Results", output_data), "predictor.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.error(error_message)

    if clear:
        st.session_state.results_history = []

    if st.session_state.results_history:
        st.subheader("History")
        st.table(pd.DataFrame(st.session_state.results_history))

# Footer
st.markdown("---")
st.markdown("Developed for Njala University | Powered by Streamlit")
