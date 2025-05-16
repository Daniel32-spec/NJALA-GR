import streamlit as st
import pandas as pd
import joblib
import re
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize session state
if 'results_history' not in st.session_state:
    st.session_state.results_history = []

# Load and cache model
def load_model():
    try:
        df = pd.read_csv('njala_student_data.csv')
        df['Result'] = df['Result'].map({'Pass':1,'Fail':0})
        X = df[['GPA','Credit_Hours','Year_Average']]
        y = df['Result']
        from sklearn.ensemble import RandomForestClassifier
        model = RandomForestClassifier(random_state=42)
        model.fit(X, y)
        return model
    except Exception as e:
        st.error(f"Error loading model: {e}")
        return None

model = load_model()

# Utility functions

def score_to_letter_grade(score):
    score = float(score)
    if score >= 75:
        return 'A'
    elif score >= 65:
        return 'B'
    elif score >= 50:
        return 'C'
    elif score >= 40:
        return 'D'
    elif score >= 30:
        return 'E'
    else:
        return 'F'

def score_to_pass_fail(score):
    return 'Pass' if float(score) >= 50 else 'Fail'

def grade_to_gpa(grade):
    mapping = {'A':4.0,'B':3.0,'C':2.0,'D':1.0,'E':0.7,'F':0.0}
    try:
        val = float(grade)
        letter = score_to_letter_grade(val)
        return mapping[letter]
    except:
        return mapping.get(str(grade).upper(), None)

# Output functions
def create_word_doc(title, data):
    doc = Document()
    doc.add_heading(title, 0)
    table = doc.add_table(rows=data.shape[0]+1, cols=len(data.columns))
    table.style = 'Table Grid'
    for j, col in enumerate(data.columns):
        cell = table.cell(0, j)
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in data.iterrows():
        for j, val in enumerate(row):
            table.cell(i+1, j).text = str(val)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def create_excel_file(data):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        data.to_excel(writer, index=False, sheet_name='Results')
    return output.getvalue()

# Predictor UI
st.subheader("Pass/Fail Predictor")
st.markdown("Enter module details or upload an Excel file to predict pass/fail, then download results as Excel or Word.")

# File uploader
uploaded_file = st.file_uploader("Upload Excel (Module, Credit_Hours, Score)", type=['xlsx','xls'])

# If user uploaded a file, process it
if uploaded_file and model:
    df = pd.read_excel(uploaded_file)
    df['Letter'] = df['Score'].apply(score_to_letter_grade)
    df['GPA'] = df['Score'].apply(grade_to_gpa)
    df['Year_Average'] = df['GPA']
    df['Credit_Hours'] = df['Credit_Hours']
    preds = model.predict(df[['GPA','Credit_Hours','Year_Average']])
    df['Result'] = ['Pass' if p==1 else 'Fail' for p in preds]
    st.dataframe(df)
    fmt = st.selectbox("Download as", ['Excel','Word'], key="upload_fmt")
    if fmt=='Excel':
        st.download_button("Download XLSX", create_excel_file(df), file_name='results.xlsx')
    else:
        st.download_button("Download DOCX", create_word_doc('Results', df), file_name='results.docx')
else:
    # Manual input if no file
    num_modules = st.number_input("Number of Modules", min_value=1, max_value=10, value=1, step=1)
    module_info = []
    total_credits = 0
    total_score = 0
    for i in range(num_modules):
        cols = st.columns(3)
        name = cols[0].text_input(f"Module {i+1} Name", key=f"mod_name_{i}")
        credits = cols[1].number_input(f"Credit Hours", 1, 6, 3, key=f"credit_{i}")
        score = cols[2].number_input(f"Score (0-100)", 0.0, 100.0, 50.0, 0.1, key=f"score_{i}")
        letter = score_to_letter_grade(score)
        gpa_val = grade_to_gpa(score)
        module_info.append({'Module': name, 'Credits': credits, 'Score': score, 'Letter': letter, 'GPA': gpa_val})
        total_credits += credits
        total_score += score * credits
    # Compute overall metrics
    weighted_avg_score = total_score / total_credits if total_credits else 0
    overall_letter = score_to_letter_grade(weighted_avg_score)
    overall_gpa = grade_to_gpa(weighted_avg_score)
    if st.button("Predict Overall") and model:
        prediction = model.predict([[overall_gpa, total_credits, overall_gpa]])[0]
        result = 'Pass' if prediction==1 else 'Fail'
        st.success(f"Overall Result: {result}")
        output_df = pd.DataFrame(module_info)
        output_df['Overall Score'] = weighted_avg_score
        output_df['Overall Letter'] = overall_letter
        output_df['Result'] = result
        fmt = st.selectbox("Download as", ['Excel','Word'], key="manual_fmt")
        if fmt == 'Excel':
            st.download_button("Download XLSX", create_excel_file(output_df), file_name='predictor_results.xlsx')
        else:
            st.download_button("Download DOCX", create_word_doc('Predictor Results', output_df), file_name='predictor_results.docx')

# Display history
if st.session_state.results_history:
    st.subheader("History")
    st.table(pd.DataFrame(st.session_state.results_history))
